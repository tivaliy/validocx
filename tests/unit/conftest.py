#
#    Copyright 2018 Vitalii Kulanov
#

import pytest

from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


@pytest.fixture(scope="module")
def document():
    doc = Document()
    doc.add_heading('Fake Title', 0)
    doc.add_heading('Fake Header 1', 1)
    doc.add_heading('Fake Header 2', 2)
    # Redefine font parameters for 'Normal' style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    # Add some text (paragraph) of 'Normal' style
    p = doc.add_paragraph('Some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    # Set custom paragraph attributes
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.left_indent = Cm(0.5)
    paragraph_format.right_indent = Cm(0.5)
    paragraph_format.space_before = Cm(1)
    paragraph_format.space_after = Cm(1)
    paragraph_format.first_line_indent = Cm(1.25)
    paragraph_format.line_spacing = Cm(1.0)
    paragraph_format.keep_with_next = True
    # Add new section
    doc.add_section()
    return doc
