#
#    Copyright 2018 Vitalii Kulanov
#

import pytest

from docx import Document
from docx.shared import Pt

from validocx import wrapper


@pytest.fixture
def document():
    doc = Document()
    doc.add_heading('Fake Title', 0)
    doc.add_heading('Fake Header 1', 1)
    doc.add_heading('Fake Header 2', 2)
    # Redefine font parameters for 'Normal' style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    # Add some text in 'Normal' style
    p = doc.add_paragraph('Some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    return doc


@pytest.fixture
def docx_wrapper(document):
    document_wrapper = wrapper.DocumentWrapper(document)
    return document_wrapper


def test_fetch_author_data(docx_wrapper):
    assert docx_wrapper.author == 'python-docx'


def test_fetch_date_of_creation(docx_wrapper):
    assert docx_wrapper.created.isoformat(' ') == '2013-12-23 23:15:00'


def test_fetch_date_of_modification(docx_wrapper):
    assert docx_wrapper.modified.isoformat(' ') == '2013-12-23 23:15:00'


def test_fetch_last_modifier_data(docx_wrapper):
    assert docx_wrapper.last_modified_by == ''


@pytest.mark.parametrize('paragraph_styles, expected', [
    ('Title', ['Fake Title']),
    (None, ['Fake Title', 'Fake Header 1',
            'Fake Header 2', 'Some bold and some italic.'])
])
def test_fetch_paragraph_data(docx_wrapper, paragraph_styles, expected):
    for item in docx_wrapper.iter_paragraphs(paragraph_styles):
        assert item.text in expected


def test_fetch_section_data(docx_wrapper):
    assert len(list(docx_wrapper.iter_sections())) == 1


@pytest.mark.parametrize('p_style, expected', [
    ('Title', [[26.0, 'Calibri']]),
    ('Normal', [[12.0, 'Calibri'], [12.0, 'Calibri', 'bold'],
                [12.0, 'Calibri'], [12.0, 'Calibri', 'italic']])
])
def test_fetch_font_attributes(docx_wrapper, p_style, expected):
    p = list(docx_wrapper.iter_paragraphs(p_style))[0]
    runs = docx_wrapper.get_font_attributes(p)
    assert runs == expected
